import docx
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
import re
import os

def main():
    doc = docx.Document()
    
    file_name = '三维课程设计期末考核与4C大赛作品要求.md'
    out_name = '三维课程设计期末考核与4C大赛作品要求.docx'
    
    with open(file_name, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip('\n') 
        if not line.strip():
            continue
            
        level = 0
        if line.startswith('# '):
            level = 1
            text = line[2:]
        elif line.startswith('## '):
            level = 2
            text = line[3:]
        elif line.startswith('### '):
            level = 3
            text = line[4:]
        else:
            text = line
            
        if level > 0:
            p = doc.add_paragraph(style=f'Heading {level}')
        else:
            if text.lstrip().startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                text = text.lstrip()[2:]
            elif text.lstrip().startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                text = text.lstrip()[2:]
            else:
                p = doc.add_paragraph()
                
        # Split text by bold marker: ** ... **
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run_text = part[2:-2].replace('\\', '').replace('`', '')
                if not run_text: continue
                run = p.add_run(run_text)
                run.bold = True
                run.font.color.rgb = RGBColor(220, 38, 38) # Red for emphasis
            else:
                run_text = part.replace('\\', '').replace('`', '')
                if not run_text: continue
                run = p.add_run(run_text)
                run.font.color.rgb = RGBColor(0, 0, 0) # Black
                
            # Set font to nicely support Chinese
            run.font.name = 'Microsoft YaHei'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
    doc.save(out_name)

if __name__ == '__main__':
    main()
